from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import pandas as pd
from copy import deepcopy
from massive_docx.util import delete_paragraph


class Base:

    def _show_styles(self):
        for style in self.template.styles:
            print(style.name)

    def _parse_template(self):
        self.template = Document(self.template_name)

    def _parse_mailing_list(self):
        self.mailing_list = pd.read_excel(self.mailing_list_name)


class Template(Base):

    def __init__(self, mailing_list_name=None, template_name=None):
        self.mailing_list_name = mailing_list_name
        self.template_name = template_name

    def process(self):
        if self.mailing_list_name is None:
            print('Please provide mailing_list_name argument first')
            raise AttributeError

        if self.template_name is None:
            print('Please provide template_name argument first')
            raise AttributeError

        target_document = Document()

        self._parse_mailing_list()

        for index, row in self.mailing_list.iterrows():

            self._parse_template()

            for paragraph in self.template.paragraphs:
                if '#Company#' in paragraph.text:
                    # print(paragraph.text)
                    paragraph.text = paragraph.text.replace('#Company#',
                                                            '' if pd.isnull(row['Company']) else str(row['Company']))
                elif '#Address#' in paragraph.text:
                    # print(paragraph.text)
                    if pd.isnull(row['Address']):
                        delete_paragraph(paragraph)
                    else:
                        paragraph.text = paragraph.text.replace('#Address#', str(row['Address']))
                elif '#City#' in paragraph.text:
                    # print(paragraph.text)
                    paragraph.text = paragraph.text.replace('#City#',
                                                            '' if pd.isnull(row['City']) else str(row['City']))
                elif '#Country#' in paragraph.text:
                    # print(paragraph.text)
                    paragraph.text = paragraph.text.replace('#Country#',
                                                            '' if pd.isnull(row['Country']) else str(row['Country']))

            self.template.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

            for element in self.template.element.body:
                new_element = deepcopy(element)
                target_document.element.body.append(new_element)

            for table in self.template.tables:
                new_table = deepcopy(table)
                # new_table.style = table.style
                target_document.tables.append(new_table)

        return target_document
